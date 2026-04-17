"""
Generate AOP Deployment Guide as a Word .docx document.
Run: python docs/generate_docx.py
Output: docs/AOP_Deployment_Guide.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

OUTPUT = "docs/AOP_Deployment_Guide.docx"

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.54)
    section.right_margin  = Cm(2.54)

# ── Colour palette ────────────────────────────────────────────────────────────
NAVY   = RGBColor(0x1A, 0x2B, 0x4A)   # aop-navy
GOLD   = RGBColor(0xC9, 0x9A, 0x06)   # aop-gold
DARK   = RGBColor(0x1F, 0x2D, 0x3D)   # aop-dark
GRAY   = RGBColor(0x4B, 0x55, 0x63)
CODE_BG = RGBColor(0xF3, 0xF4, 0xF6)
CODE_FG = RGBColor(0x1F, 0x29, 0x37)

# ── Helpers ───────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color="F3F4F6"):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def heading1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.font.size  = Pt(18)
    run.font.bold  = True
    run.font.color.rgb = NAVY
    # bottom border
    pPr = p._p.get_or_add_pPr()
    pb  = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "C99A06")
    pb.append(bot)
    pPr.append(pb)
    return p

def heading2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.font.size  = Pt(13)
    run.font.bold  = True
    run.font.color.rgb = DARK
    return p

def heading3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.font.size  = Pt(11)
    run.font.bold  = True
    run.font.color.rgb = GRAY
    return p

def body(text, bold=False, italic=False, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size   = Pt(10)
    run.font.bold   = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    return p

def note(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run("ℹ  " + text)
    run.font.size   = Pt(9.5)
    run.font.italic = True
    run.font.color.rgb = GOLD
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent  = Cm(0.5 + level * 0.5)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(10)
    return p

def checkbox(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("☐  " + text)
    run.font.size = Pt(10)
    return p

def code_block(lines: list[str]):
    """Render lines as a shaded code block (one table row per block)."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, "1F2937")
    cell._tc.get_or_add_tcPr()
    # clear default paragraph
    cell.paragraphs[0].clear()
    for i, line in enumerate(lines):
        if i == 0:
            p = cell.paragraphs[0]
        else:
            p = cell.add_paragraph()
        run = p.add_run(line)
        run.font.name  = "Courier New"
        run.font.size  = Pt(8.5)
        run.font.color.rgb = RGBColor(0xA7, 0xF3, 0xD0)   # mint green
    doc.add_paragraph()   # spacing after block

def table_header_row(table, headers):
    row = table.rows[0]
    for i, h in enumerate(headers):
        cell = row.cells[i]
        set_cell_bg(cell, "1A2B4A")
        run = cell.paragraphs[0].add_run(h)
        run.font.bold  = True
        run.font.size  = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

def add_table_row(table, values, shade=False):
    row = table.add_row()
    for i, v in enumerate(values):
        cell = row.cells[i]
        if shade:
            set_cell_bg(cell, "F9FAFB")
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(str(v))
        run.font.size = Pt(9)
        run.font.name = "Courier New" if i > 0 else run.font.name
    return row

# ══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(60)
p.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("AOP Platform")
run.font.size  = Pt(32)
run.font.bold  = True
run.font.color.rgb = NAVY

p2 = doc.add_paragraph()
p2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run("Deployment Guide — Go Live")
run2.font.size  = Pt(18)
run2.font.color.rgb = GOLD

doc.add_paragraph()
p3 = doc.add_paragraph()
p3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run("17-Step Complete Setup · us-east-1 (N. Virginia)")
run3.font.size  = Pt(11)
run3.font.color.rgb = GRAY

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (manual)
# ══════════════════════════════════════════════════════════════════════════════
heading1("Table of Contents")
toc_items = [
    ("1",  "Prerequisites"),
    ("2",  "Bootstrap Terraform State Backend"),
    ("3",  "Create ECR Repositories"),
    ("4",  "Register Your Domain & Route 53 Hosted Zone"),
    ("5",  "Configure Terraform Variables"),
    ("6",  "Provision Infrastructure with Terraform"),
    ("7",  "Build & Push Docker Images"),
    ("8",  "Configure Application Secrets in AWS"),
    ("9",  "Set Up AWS SES (Email)"),
    ("10", "Configure GitHub Secrets"),
    ("11", "Configure GitHub Environments"),
    ("12", "Set Up OIDC Trust Between GitHub and AWS"),
    ("13", "Run the First Database Migration & Seed"),
    ("14", "Deploy Staging"),
    ("15", "Deploy Production"),
    ("16", "DNS Cutover"),
    ("17", "Post-Go-Live Checklist"),
]
for num, title in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(f"  {num}.  {title}")
    run.font.size = Pt(10)
    run.font.color.rgb = NAVY

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# STEP 1 — PREREQUISITES
# ══════════════════════════════════════════════════════════════════════════════
heading1("Step 1 — Prerequisites")
body("Install and verify the following tools before starting. All commands below run in a terminal (bash / zsh / PowerShell with WSL).")

heading2("Required Tools")
tbl = doc.add_table(rows=1, cols=3)
tbl.style = "Table Grid"
table_header_row(tbl, ["Tool", "Minimum Version", "Install"])
tool_rows = [
    ("AWS CLI",    "v2",   "brew install awscli  OR  https://aws.amazon.com/cli"),
    ("Terraform",  "1.6+", "brew install terraform  OR  https://developer.hashicorp.com/terraform/downloads"),
    ("Docker",     "24+",  "https://docs.docker.com/get-docker/"),
    ("pnpm",       "8+",   "npm install -g pnpm"),
    ("GitHub CLI", "2+",   "brew install gh  OR  https://cli.github.com"),
    ("jq",         "any",  "brew install jq  OR  https://jqlang.github.io/jq/"),
]
for i, (tool, ver, inst) in enumerate(tool_rows):
    add_table_row(tbl, [tool, ver, inst], shade=(i % 2 == 1))
doc.add_paragraph()

heading2("AWS Account Setup")
for step in [
    "1. Log in to console.aws.amazon.com",
    "2. Go to IAM → Users → Create user  (or use an existing admin user)",
    "3. Attach the AdministratorAccess managed policy  (scope down after go-live)",
    "4. Under the user → Security credentials → Create access key → choose 'CLI'",
    "5. Copy the Access Key ID and Secret Access Key",
]:
    body(step)

heading3("Configure your terminal:")
code_block([
    "aws configure",
    "  AWS Access Key ID:     AKIA...",
    "  AWS Secret Access Key: <your secret>",
    "  Default region:        us-east-1",
    "  Default output format: json",
    "",
    "# Verify",
    "aws sts get-caller-identity",
])

note("us-east-1 (N. Virginia) is available by default — no opt-in required.")

# ══════════════════════════════════════════════════════════════════════════════
# STEP 2 — BOOTSTRAP TERRAFORM BACKEND
# ══════════════════════════════════════════════════════════════════════════════
heading1("Step 2 — Bootstrap Terraform State Backend")
body("Terraform stores state in S3 and uses DynamoDB for locking. These resources must exist before running terraform init. Create them once — never delete them.")

heading2("Create S3 State Bucket + DynamoDB Locks Table")
code_block([
    "# 1. Create the bucket",
    "aws s3 mb s3://aop-terraform-state --region us-east-1",
    "",
    "# 2. Enable versioning",
    "aws s3api put-bucket-versioning \\",
    "  --bucket aop-terraform-state \\",
    "  --versioning-configuration Status=Enabled \\",
    "  --region us-east-1",
    "",
    "# 3. Enable KMS encryption",
    "aws s3api put-bucket-encryption \\",
    "  --bucket aop-terraform-state \\",
    "  --server-side-encryption-configuration '{",
    '    "Rules": [{"ApplyServerSideEncryptionByDefault":{"SSEAlgorithm":"aws:kms"}}]',
    "  }' --region us-east-1",
    "",
    "# 4. Block all public access",
    "aws s3api put-public-access-block \\",
    "  --bucket aop-terraform-state \\",
    "  --public-access-block-configuration \\",
    '    "BlockPublicAcls=true,IgnorePublicAcls=true,BlockPublicPolicy=true,RestrictPublicBuckets=true" \\',
    "  --region us-east-1",
    "",
    "# 5. Create DynamoDB lock table",
    "aws dynamodb create-table \\",
    "  --table-name aop-terraform-locks \\",
    "  --attribute-definitions AttributeName=LockID,AttributeType=S \\",
    "  --key-schema AttributeName=LockID,KeyType=HASH \\",
    "  --billing-mode PAY_PER_REQUEST \\",
    "  --region us-east-1",
    "",
    "# Verify",
    'aws dynamodb describe-table --table-name aop-terraform-locks --region us-east-1 --query Table.TableStatus',
    "# Expected: \"ACTIVE\"",
])

# ══════════════════════════════════════════════════════════════════════════════
# STEP 3 — ECR REPOSITORIES
# ══════════════════════════════════════════════════════════════════════════════
heading1("Step 3 — Create ECR Repositories")
body("Three ECR repositories are needed — one per service (API, Web, Worker). ECR stores the Docker images that ECS pulls at deploy time.")

code_block([
    "export AWS_ACCOUNT_ID=$(aws sts get-caller-identity --query Account --output text)",
    'export AWS_REGION="us-east-1"',
    'export ECR_BASE="${AWS_ACCOUNT_ID}.dkr.ecr.${AWS_REGION}.amazonaws.com"',
    "",
    "for REPO in aop-api aop-web aop-worker; do",
    '  aws ecr create-repository \\',
    '    --repository-name "$REPO" \\',
    '    --image-scanning-configuration scanOnPush=true \\',
    '    --encryption-configuration encryptionType=KMS \\',
    '    --region "$AWS_REGION" 2>/dev/null \\',
    '    && echo "Created $REPO" || echo "WARNING: $REPO already exists"',
    "done",
    "",
    'echo "ECR base URI: ${ECR_BASE}"',
    "# Write this down — used in Steps 5 and 7",
])

# ══════════════════════════════════════════════════════════════════════════════
# STEP 4 — ROUTE 53
# ══════════════════════════════════════════════════════════════════════════════
heading1("Step 4 — Register Domain & Route 53 Hosted Zone")

heading2("Option A — Domain registered elsewhere (GoDaddy, Namecheap, etc.)")
for s in [
    "1. Create a hosted zone in Route 53:",
]:
    body(s)
code_block([
    "aws route53 create-hosted-zone \\",
    '  --name "yourdomain.com" \\',
    '  --caller-reference "$(date +%s)"',
])
body("2. The response contains 4 NS records (ns-xxx.awsdns-xx.com). Copy all four.")
body("3. Log in to your registrar and replace the existing nameservers with those 4 values.")
body("4. Wait 15–60 minutes for DNS propagation.")

heading2("Option B — Register domain via Route 53")
body("AWS Console → Route 53 → Domains → Register domain")

heading2("Get Your Zone ID")
code_block([
    "aws route53 list-hosted-zones \\",
    "  --query \"HostedZones[?Name=='yourdomain.com.'].Id\" \\",
    "  --output text",
    "# Returns: /hostedzone/Z1234567890ABCDEF",
    "# Your Zone ID is: Z1234567890ABCDEF  (strip the prefix)",
])
note("Save the Zone ID — it is required in Step 5 (tfvars) and Step 16 (DNS cutover).")

# ══════════════════════════════════════════════════════════════════════════════
# STEP 5 — TERRAFORM VARIABLES
# ══════════════════════════════════════════════════════════════════════════════
heading1("Step 5 — Configure Terraform Variables")
body("Create variable files for each environment. Replace all placeholder values.")

code_block([
    "cd infra",
    "cp dev.tfvars.example staging.tfvars",
    "cp dev.tfvars.example production.tfvars",
])

heading2("infra/staging.tfvars")
code_block([
    'aws_region   = "us-east-1"',
    'project_name = "aop"',
    'vpc_cidr     = "10.0.0.0/16"',
    "",
    'domain_name     = "staging.yourdomain.com"   # your staging subdomain',
    'route53_zone_id = "Z1234567890ABCDEF"         # from Step 4',
    'alert_email     = "yourteam@yourdomain.com"',
    "",
    "# Replace 123456789012 with your AWS_ACCOUNT_ID",
    'api_container_image    = "123456789012.dkr.ecr.us-east-1.amazonaws.com/aop-api:latest"',
    'worker_container_image = "123456789012.dkr.ecr.us-east-1.amazonaws.com/aop-worker:latest"',
    "",
    "api_cpu    = 512    api_memory    = 1024",
    "worker_cpu = 256    worker_memory = 512",
    '',
    'db_instance_class       = "db.t3.medium"',
    "db_allocated_storage_gb = 20",
    'redis_node_type         = "cache.t3.micro"',
])

heading2("infra/production.tfvars")
code_block([
    'aws_region   = "us-east-1"',
    'project_name = "aop"',
    'vpc_cidr     = "10.1.0.0/16"',
    "",
    'domain_name     = "yourdomain.com"            # apex domain',
    'route53_zone_id = "Z1234567890ABCDEF"',
    'alert_email     = "yourteam@yourdomain.com"',
    "",
    'api_container_image    = "123456789012.dkr.ecr.us-east-1.amazonaws.com/aop-api:latest"',
    'worker_container_image = "123456789012.dkr.ecr.us-east-1.amazonaws.com/aop-worker:latest"',
    "",
    "api_cpu    = 1024   api_memory    = 2048",
    "worker_cpu = 512    worker_memory = 1024",
    "",
    'db_instance_class       = "db.r6g.large"',
    "db_allocated_storage_gb = 100",
    'redis_node_type         = "cache.r6g.large"',
])

heading2("Keep tfvars out of git")
code_block([
    'echo "infra/staging.tfvars"    >> .gitignore',
    'echo "infra/production.tfvars" >> .gitignore',
])

# ══════════════════════════════════════════════════════════════════════════════
# STEP 6 — PROVISION INFRASTRUCTURE
# ══════════════════════════════════════════════════════════════════════════════
heading1("Step 6 — Provision Infrastructure with Terraform")
body("This creates the entire AWS infrastructure: VPC, subnets, RDS, ElastiCache, ECS cluster, ALB, CloudFront, WAF, S3 buckets, KMS keys, IAM roles, CloudWatch. Takes 15–25 minutes per environment.")

code_block([
    "cd infra",
    "terraform init          # downloads providers, connects to S3 backend",
    "",
    "# ── STAGING ──────────────────────────────────────────────────",
    "terraform workspace new staging",
    "terraform workspace select staging",
    "",
    "terraform plan -var-file=staging.tfvars -out=staging.plan",
    "# Review the plan (~50–70 resources). When satisfied:",
    "terraform apply staging.plan",
    "",
    "# Capture outputs",
    "terraform output -json > ../infra-outputs-staging.json",
    "export STAGING_ECS_CLUSTER=$(terraform output -raw ecs_cluster_name)",
    "export STAGING_CF_DOMAIN=$(terraform output -raw cloudfront_domain)",
    "export STAGING_DB_SECRET_ARN=$(terraform output -raw db_secret_arn)",
    "",
    "# ── PRODUCTION ───────────────────────────────────────────────",
    "terraform workspace new production",
    "terraform workspace select production",
    "",
    "terraform plan -var-file=production.tfvars -out=production.plan",
    "terraform apply production.plan",
    "",
    "terraform output -json > ../infra-outputs-production.json",
    "export PROD_ECS_CLUSTER=$(terraform output -raw ecs_cluster_name)",
    "export PROD_CF_DOMAIN=$(terraform output -raw cloudfront_domain)",
    "export PROD_DB_SECRET_ARN=$(terraform output -raw db_secret_arn)",
])
note("If apply fails — read the error carefully. Common causes: IAM permission missing, quota limit hit. Fix the root cause; never force-delete Terraform state.")

# ══════════════════════════════════════════════════════════════════════════════
# STEP 7 — BUILD & PUSH DOCKER IMAGES
# ══════════════════════════════════════════════════════════════════════════════
heading1("Step 7 — Build & Push Docker Images")
body("Seed ECR with the initial images so ECS has something to run. CI/CD will handle subsequent image updates automatically.")

code_block([
    "cd /path/to/aop    # repo root",
    "",
    "# 1. Authenticate Docker to ECR",
    "aws ecr get-login-password --region us-east-1 | \\",
    '  docker login --username AWS --password-stdin "${ECR_BASE}"',
    "",
    "# 2. Build all three images",
    "#    Add --platform linux/amd64 if building on Apple Silicon",
    "docker build -f docker/api.Dockerfile \\",
    '  -t "${ECR_BASE}/aop-api:latest" \\',
    '  -t "${ECR_BASE}/aop-api:initial" \\',
    "  --platform linux/amd64 .",
    "",
    "docker build -f docker/web.Dockerfile \\",
    '  -t "${ECR_BASE}/aop-web:latest" \\',
    '  -t "${ECR_BASE}/aop-web:initial" \\',
    "  --platform linux/amd64 .",
    "",
    "docker build -f docker/worker.Dockerfile \\",
    '  -t "${ECR_BASE}/aop-worker:latest" \\',
    '  -t "${ECR_BASE}/aop-worker:initial" \\',
    "  --platform linux/amd64 .",
    "",
    "# 3. Push to ECR",
    'docker push "${ECR_BASE}/aop-api:latest"    && docker push "${ECR_BASE}/aop-api:initial"',
    'docker push "${ECR_BASE}/aop-web:latest"    && docker push "${ECR_BASE}/aop-web:initial"',
    'docker push "${ECR_BASE}/aop-worker:latest" && docker push "${ECR_BASE}/aop-worker:initial"',
    "",
    "# 4. Re-apply Terraform so ECS picks up the real images",
    "cd infra && terraform workspace select staging",
    "terraform apply -var-file=staging.tfvars",
])

# ══════════════════════════════════════════════════════════════════════════════
# STEP 8 — SECRETS MANAGER
# ══════════════════════════════════════════════════════════════════════════════
heading1("Step 8 — Configure Application Secrets in AWS")
body("The app reads credentials from AWS Secrets Manager at runtime. The database secret is created automatically by Terraform. Create the remaining secrets manually for both environments.")

code_block([
    'ENV="staging"   # repeat with ENV="production" for prod',
    "",
    "# JWT signing secret",
    "JWT_SECRET=$(openssl rand -base64 48)",
    "aws secretsmanager create-secret \\",
    '  --name "aop/${ENV}/jwt-secret" \\',
    '  --secret-string "${JWT_SECRET}" \\',
    "  --region us-east-1",
    "",
    "# metals.dev gold price API key  (sign up at https://metals.dev)",
    "aws secretsmanager create-secret \\",
    '  --name "aop/${ENV}/metals-dev-api-key" \\',
    "  --secret-string '{\"METALS_DEV_API_KEY\":\"your-key-from-metals.dev\"}' \\",
    "  --region us-east-1",
    "",
    "# SMS credentials  (Africa's Talking / Twilio)",
    "aws secretsmanager create-secret \\",
    '  --name "aop/${ENV}/sms-credentials" \\',
    '  --secret-string \'{"SMS_API_KEY":"your-key","SMS_SENDER_ID":"AOP"}\' \\',
    "  --region us-east-1",
    "",
    "# Sanctions/AML screening API key",
    "aws secretsmanager create-secret \\",
    '  --name "aop/${ENV}/sanctions-api-key" \\',
    "  --secret-string '{\"SANCTIONS_API_KEY\":\"your-key\"}' \\",
    "  --region us-east-1",
])
note("Generate a separate JWT_SECRET for production using a fresh `openssl rand -base64 48` call — never reuse the staging secret.")

# ══════════════════════════════════════════════════════════════════════════════
# STEP 9 — SES
# ══════════════════════════════════════════════════════════════════════════════
heading1("Step 9 — Set Up AWS SES (Email)")
body("SES sends KYC notifications, disbursement alerts, and compliance reports.")

heading2("Verify your sending domain")
code_block([
    "aws ses verify-domain-identity \\",
    '  --domain "yourdomain.com" \\',
    "  --region us-east-1",
    "# Returns a VerificationToken — add it as a TXT record in Route 53:",
    "#   Name:  _amazonses.yourdomain.com",
    "#   Type:  TXT",
    "#   Value: <the returned token>",
    "",
    "# Check verification status (takes a few minutes)",
    "aws ses get-identity-verification-attributes \\",
    '  --identities "yourdomain.com" \\',
    "  --region us-east-1",
    '# VerificationStatus should become "Success"',
])

heading2("Request Production Access (remove sandbox)")
body("New AWS accounts can only send to verified addresses. To send to anyone:")
for s in [
    "1. Go to AWS Console → SES → Account dashboard",
    "2. Click Request production access",
    "3. Fill in the form: Mail type = Transactional, use case = KYC/compliance notifications",
    "4. AWS approves within 24 hours",
]:
    body(s)
note("Until approved, add test recipient email addresses via: aws ses verify-email-identity --email-address you@yourdomain.com --region us-east-1")

# ══════════════════════════════════════════════════════════════════════════════
# STEP 10 — GITHUB SECRETS
# ══════════════════════════════════════════════════════════════════════════════
heading1("Step 10 — Configure GitHub Secrets")
body("All secrets live at: GitHub repo → Settings → Secrets and variables → Actions")
body("Set them via the CLI (faster) or the web UI:")
code_block([
    "gh auth login   # if not already authenticated",
    'gh secret set SECRET_NAME --body "value"',
])

heading2("Shared Secrets")
tbl = doc.add_table(rows=1, cols=2)
tbl.style = "Table Grid"
table_header_row(tbl, ["Secret Name", "Value"])
shared = [
    ("AWS_REGION",               "us-east-1"),
    ("AWS_CI_ROLE_ARN",          "Set in Step 12"),
    ("AWS_DEPLOY_ROLE_ARN",      "Set in Step 12"),
    ("AWS_PROD_DEPLOY_ROLE_ARN", "Set in Step 12"),
    ("ECR_API_REPO",             "aop-api"),
    ("ECR_WEB_REPO",             "aop-web"),
    ("ECR_WORKER_REPO",          "aop-worker"),
    ("CODECOV_TOKEN",            "From codecov.io after linking repo"),
    ("SLACK_DEPLOY_WEBHOOK",     "Slack incoming webhook URL"),
]
for i, (k, v) in enumerate(shared):
    add_table_row(tbl, [k, v], shade=(i % 2 == 1))
doc.add_paragraph()

heading2("Staging Secrets")
tbl2 = doc.add_table(rows=1, cols=3)
tbl2.style = "Table Grid"
table_header_row(tbl2, ["Secret Name", "Example Value", "How to Get"])
staging_s = [
    ("STAGING_ECS_CLUSTER",      "aop-staging",               "terraform output ecs_cluster_name  (staging workspace)"),
    ("STAGING_API_SERVICE",      "aop-staging-api",            "aws ecs list-services --cluster <name>"),
    ("STAGING_WEB_SERVICE",      "aop-staging-web",            "aws ecs list-services --cluster <name>"),
    ("STAGING_WORKER_SERVICE",   "aop-staging-worker",         "aws ecs list-services --cluster <name>"),
    ("STAGING_URL",              "https://staging.yourdomain.com", "Your staging domain"),
    ("STAGING_RDS_INSTANCE_ID",  "aop-staging-db",             "AWS Console → RDS → DB identifier"),
    ("STAGING_MIGRATION_TASK_DEF", "arn:aws:ecs:...",          "aws ecs list-task-definitions --family-prefix aop-staging-migrate"),
    ("STAGING_NETWORK_CONFIG",   "{\"awsvpcConfiguration\":{...}}", "See Step 10 in DEPLOYMENT.md"),
]
for i, row in enumerate(staging_s):
    add_table_row(tbl2, row, shade=(i % 2 == 1))
doc.add_paragraph()

heading2("Production Secrets")
tbl3 = doc.add_table(rows=1, cols=2)
tbl3.style = "Table Grid"
table_header_row(tbl3, ["Secret Name", "Value"])
prod_s = [
    ("PROD_ECS_CLUSTER",        "aop-production"),
    ("PROD_API_SERVICE",        "aop-production-api"),
    ("PROD_WEB_SERVICE",        "aop-production-web"),
    ("PROD_WORKER_SERVICE",     "aop-production-worker"),
    ("PROD_URL",                "https://yourdomain.com"),
    ("PROD_RDS_INSTANCE_ID",    "aop-production-db"),
    ("PROD_MIGRATION_TASK_DEF", "arn:aws:ecs:..."),
    ("PROD_NETWORK_CONFIG",     "{\"awsvpcConfiguration\":{...}}"),
]
for i, (k, v) in enumerate(prod_s):
    add_table_row(tbl3, [k, v], shade=(i % 2 == 1))
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# STEP 11 — GITHUB ENVIRONMENTS
# ══════════════════════════════════════════════════════════════════════════════
heading1("Step 11 — Configure GitHub Environments")
body("Environments gate which branch can deploy where and add a human approval step for production.")

heading2("Create staging environment (auto-deploys from main)")
code_block([
    "gh api repos/Auzyem/aop/environments/staging \\",
    "  --method PUT \\",
    "  --field deployment_branch_policy='{\"protected_branches\":false,\"custom_branch_policies\":true}'",
    "",
    "gh api repos/Auzyem/aop/environments/staging/deployment-branch-policies \\",
    "  --method POST \\",
    "  --field name=main",
])

heading2("Create production environment (requires manual approval)")
code_block([
    "gh api repos/Auzyem/aop/environments/production \\",
    "  --method PUT \\",
    "  --field deployment_branch_policy='{\"protected_branches\":true,\"custom_branch_policies\":false}'",
])

heading2("Add required reviewers (must be done in the web UI)")
for s in [
    "1. Go to github.com/Auzyem/aop → Settings → Environments → production",
    "2. Click Required reviewers → add yourself (and any other approvers)",
    "3. Check Prevent self-review if a second person must always approve",
    "4. Save protection rules",
]:
    body(s)

# ══════════════════════════════════════════════════════════════════════════════
# STEP 12 — OIDC TRUST
# ══════════════════════════════════════════════════════════════════════════════
heading1("Step 12 — Set Up OIDC Trust Between GitHub and AWS")
body("OIDC allows GitHub Actions to assume AWS IAM roles without storing long-lived credentials as GitHub secrets.")

heading2("1. Register GitHub as an OIDC provider (once per AWS account)")
code_block([
    "aws iam create-open-id-connect-provider \\",
    "  --url https://token.actions.githubusercontent.com \\",
    "  --client-id-list sts.amazonaws.com \\",
    "  --thumbprint-list 6938fd4d98bab03faadb97b34396831e3780aea1",
])

heading2("2. Create the three IAM roles")
code_block([
    'export GITHUB_ORG="Blaqlv"',
    'export REPO="aop"',
    "export ACCOUNT_ID=$(aws sts get-caller-identity --query Account --output text)",
    'export OIDC_ARN="arn:aws:iam::${ACCOUNT_ID}:oidc-provider/token.actions.githubusercontent.com"',
    "",
    "# ── CI Role (ECR push from any branch) ───────────────────────",
    "cat > /tmp/ci-trust.json << EOF",
    "{",
    '  "Version": "2012-10-17",',
    '  "Statement": [{"Effect":"Allow",',
    '    "Principal":{"Federated":"${OIDC_ARN}"},',
    '    "Action":"sts:AssumeRoleWithWebIdentity",',
    '    "Condition":{"StringEquals":',
    '      {"token.actions.githubusercontent.com:aud":"sts.amazonaws.com"},',
    '     "StringLike":',
    '      {"token.actions.githubusercontent.com:sub":"repo:${GITHUB_ORG}/${REPO}:*"}',
    "    }",
    "  }]",
    "}",
    "EOF",
    "aws iam create-role --role-name aop-github-ci \\",
    "  --assume-role-policy-document file:///tmp/ci-trust.json",
    "aws iam attach-role-policy --role-name aop-github-ci \\",
    "  --policy-arn arn:aws:iam::aws:policy/AmazonEC2ContainerRegistryPowerUser",
    "",
    "# ── Staging Deploy Role ──────────────────────────────────────",
    "# (trust condition: environment:staging)",
    "aws iam create-role --role-name aop-github-deploy-staging \\",
    "  --assume-role-policy-document file:///tmp/staging-trust.json",
    "aws iam attach-role-policy --role-name aop-github-deploy-staging \\",
    "  --policy-arn arn:aws:iam::aws:policy/AmazonECS_FullAccess",
    "",
    "# ── Production Deploy Role ───────────────────────────────────",
    "# (trust condition: environment:production)",
    "aws iam create-role --role-name aop-github-deploy-production \\",
    "  --assume-role-policy-document file:///tmp/prod-trust.json",
    "aws iam attach-role-policy --role-name aop-github-deploy-production \\",
    "  --policy-arn arn:aws:iam::aws:policy/AmazonECS_FullAccess",
])

heading2("3. Save role ARNs as GitHub secrets")
code_block([
    "CI_ROLE=$(aws iam get-role --role-name aop-github-ci --query Role.Arn --output text)",
    "STAGING_ROLE=$(aws iam get-role --role-name aop-github-deploy-staging --query Role.Arn --output text)",
    "PROD_ROLE=$(aws iam get-role --role-name aop-github-deploy-production --query Role.Arn --output text)",
    "",
    'gh secret set AWS_CI_ROLE_ARN          --body "$CI_ROLE"',
    'gh secret set AWS_DEPLOY_ROLE_ARN      --body "$STAGING_ROLE"',
    'gh secret set AWS_PROD_DEPLOY_ROLE_ARN --body "$PROD_ROLE"',
])

# ══════════════════════════════════════════════════════════════════════════════
# STEP 13 — MIGRATION
# ══════════════════════════════════════════════════════════════════════════════
heading1("Step 13 — Run the First Database Migration & Seed")
body("The database schema must be created before the app can start. Run this for both staging and production.")

heading2("Option A — ECS run-task (recommended)")
code_block([
    "TASK_ARN=$(aws ecs run-task \\",
    '  --cluster "${STAGING_ECS_CLUSTER}" \\',
    '  --task-definition "aop-staging-migrate" \\',
    "  --launch-type FARGATE \\",
    '  --network-configuration "${STAGING_NETWORK_CONFIG}" \\',
    "  --overrides '{",
    '    "containerOverrides": [{"name":"migrate",',
    '    "command":["pnpm","--filter","@aop/db","migrate:deploy"]}]',
    "  }' \\",
    "  --region us-east-1 \\",
    "  --query 'tasks[0].taskArn' --output text)",
    "",
    "# Wait for completion",
    "aws ecs wait tasks-stopped \\",
    '  --cluster "${STAGING_ECS_CLUSTER}" \\',
    '  --tasks "${TASK_ARN}" \\',
    "  --region us-east-1",
    "",
    "# Check exit code (0 = success)",
    "aws ecs describe-tasks \\",
    '  --cluster "${STAGING_ECS_CLUSTER}" \\',
    '  --tasks "${TASK_ARN}" \\',
    "  --region us-east-1 \\",
    "  --query 'tasks[0].containers[0].exitCode'",
])

heading2("Option B — Direct connection with DATABASE_URL")
code_block([
    "DB_SECRET=$(aws secretsmanager get-secret-value \\",
    '  --secret-id "${STAGING_DB_SECRET_ARN}" \\',
    "  --query SecretString --output text)",
    "",
    "DATABASE_URL=\"postgresql://$(echo $DB_SECRET | jq -r .username):$(echo $DB_SECRET | jq -r .password)@$(echo $DB_SECRET | jq -r .host):$(echo $DB_SECRET | jq -r .port)/$(echo $DB_SECRET | jq -r .dbname)\"",
    "",
    "DATABASE_URL=\"${DATABASE_URL}\" pnpm --filter @aop/db migrate:deploy",
    "DATABASE_URL=\"${DATABASE_URL}\" pnpm --filter @aop/db seed",
])
note("The DB is in a private subnet. For Option B you need a VPN or SSM port-forward to a bastion host.")

# ══════════════════════════════════════════════════════════════════════════════
# STEP 14 — DEPLOY STAGING
# ══════════════════════════════════════════════════════════════════════════════
heading1("Step 14 — Deploy Staging")
body("With infrastructure provisioned, images pushed, and the database migrated, trigger the first staging deployment.")

code_block([
    "# Trigger by pushing to main",
    "git push origin main",
    "",
    "# Monitor the workflow",
    "gh run watch",
    "",
    "# Wait for ECS services to stabilise",
    "aws ecs wait services-stable \\",
    '  --cluster "${STAGING_ECS_CLUSTER}" \\',
    "  --services aop-staging-api aop-staging-web aop-staging-worker \\",
    "  --region us-east-1",
    "",
    "# Health check smoke test",
    'curl -sf "https://staging.yourdomain.com/api/health" | jq \'.\'',
    '# Expected: { "status": "ok", "db": "connected", "redis": "connected" }',
    "",
    "# If it fails — tail the logs",
    "aws logs tail /ecs/aop-staging-api --follow --region us-east-1",
])

# ══════════════════════════════════════════════════════════════════════════════
# STEP 15 — DEPLOY PRODUCTION
# ══════════════════════════════════════════════════════════════════════════════
heading1("Step 15 — Deploy Production")
body("Production deployments require a manual approval via the GitHub Environment gate.")

heading2("1. Set all production GitHub secrets (Step 10 production section)")
heading2("2. Run the production database migration (Step 13 with PROD values)")
heading2("3. Trigger the production workflow")
code_block([
    "gh workflow run deploy.yml --ref main --field environment=production",
    "",
    "# The workflow will pause with: 'Waiting for review'",
])

heading2("4. Approve the deployment")
for s in [
    "1. Go to github.com/Auzyem/aop → Actions → the running workflow",
    "2. Click Review deployments",
    "3. Tick production",
    "4. Click Approve and deploy",
]:
    body(s)

heading2("5. Monitor")
code_block([
    "aws ecs wait services-stable \\",
    '  --cluster "${PROD_ECS_CLUSTER}" \\',
    "  --services aop-production-api aop-production-web aop-production-worker \\",
    "  --region us-east-1",
    "",
    'curl -sf "https://yourdomain.com/api/health" | jq \'.\'',
])

# ══════════════════════════════════════════════════════════════════════════════
# STEP 16 — DNS CUTOVER
# ══════════════════════════════════════════════════════════════════════════════
heading1("Step 16 — DNS Cutover")
body("Point your live domain to CloudFront — the public-facing entry point for the application.")

code_block([
    '# Z2FDTNDATAQYW2 is the fixed AWS hosted zone ID for ALL CloudFront distributions',
    "aws route53 change-resource-record-sets \\",
    '  --hosted-zone-id "Z1234567890ABCDEF" \\',
    "  --change-batch '{",
    '    "Changes": [{',
    '      "Action": "UPSERT",',
    '      "ResourceRecordSet": {',
    '        "Name": "yourdomain.com",',
    '        "Type": "A",',
    '        "AliasTarget": {',
    '          "HostedZoneId": "Z2FDTNDATAQYW2",',
    '          "DNSName": "d1234abcd.cloudfront.net",',
    '          "EvaluateTargetHealth": false',
    "        }",
    "      }",
    "    }]",
    "  }'",
])

heading2("Verify propagation")
code_block([
    "dig +short yourdomain.com @8.8.8.8     # Google DNS",
    "dig +short yourdomain.com @1.1.1.1     # Cloudflare DNS",
    "# Both should resolve to CloudFront IPs",
    "",
    "# Verify TLS",
    'curl -sv "https://yourdomain.com/api/health" 2>&1 | grep "SSL certificate"',
    "# Expected: * SSL certificate verify ok.",
])

# ══════════════════════════════════════════════════════════════════════════════
# STEP 17 — POST-GO-LIVE CHECKLIST
# ══════════════════════════════════════════════════════════════════════════════
heading1("Step 17 — Post-Go-Live Checklist")
body("Work through every item before declaring the platform live.")

heading2("Health & Connectivity")
for item in [
    "GET https://yourdomain.com/api/health returns { \"status\": \"ok\", \"db\": \"connected\", \"redis\": \"connected\" }",
    "Login flow works end-to-end",
    "Create a test client → run through KYC steps → verify status changes",
    "Create a test transaction → verify gold weight + settlement calculation",
]:
    checkbox(item)

heading2("Infrastructure")
for item in [
    "CloudWatch dashboards show healthy metrics (CPU < 80%, 5xx rate = 0)",
    "RDS automated backups enabled — check AWS Console → RDS → Maintenance & backups",
    "CloudTrail is recording in us-east-1",
    "SNS alert email confirmed — check alert_email inbox for the confirmation link",
]:
    checkbox(item)

heading2("Security")
for item in [
    "WAF is attached to the CloudFront distribution",
    "All S3 buckets have public access blocked",
    "Secrets Manager secrets are not publicly accessible",
    "No secrets committed to git: git log --all -- '*.env'",
]:
    checkbox(item)

heading2("CI/CD")
for item in [
    "Push a small change to main — verify staging auto-deploys successfully",
    "Verify the production deploy workflow pauses for manual approval",
    "Dependabot is enabled — Settings → Code security → Dependabot alerts",
]:
    checkbox(item)

heading2("Third-Party Integrations")
for item in [
    "SES production access granted (sandbox removed)",
    "Send a test transactional email — verify it arrives and is not flagged as spam",
    "metals.dev API key set in Secrets Manager — worker logs show gold price polls",
    "AM Fix (10:30 UTC) and PM Fix (15:00 UTC) BullMQ jobs are scheduled",
]:
    checkbox(item)

heading2("Monitoring")
for item in [
    "Trigger a test alarm — verify SNS email arrives",
    "CloudWatch log groups exist for all three services (/ecs/aop-production-*)",
]:
    checkbox(item)

heading2("Go-Live Sign-Off")
for item in [
    "Stakeholder smoke test complete",
    "On-call runbook shared with ops team (docs/RUNBOOK.md)",
    "Incident response contacts documented",
    "First production backup verified",
]:
    checkbox(item)

# ══════════════════════════════════════════════════════════════════════════════
# APPENDIX
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading1("Appendix — Common Commands")

heading2("Force ECS rolling restart (no new image)")
code_block([
    "aws ecs update-service \\",
    '  --cluster "${PROD_ECS_CLUSTER}" \\',
    "  --service aop-production-api \\",
    "  --force-new-deployment \\",
    "  --region us-east-1",
])

heading2("Tail ECS logs")
code_block([
    "aws logs tail /ecs/aop-production-api --follow --region us-east-1",
])

heading2("List all Secrets Manager secrets")
code_block([
    "aws secretsmanager list-secrets \\",
    "  --filter Key=name,Values=aop/ \\",
    "  --region us-east-1 \\",
    "  --query 'SecretList[*].Name'",
])

heading2("Check current Terraform workspace")
code_block([
    "cd infra && terraform workspace show",
])

# ── Save ──────────────────────────────────────────────────────────────────────
doc.save(OUTPUT)
print(f"Saved: {OUTPUT}")
